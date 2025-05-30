from flask import Flask, render_template, request, redirect, url_for, flash, send_from_directory, session, jsonify
import os
import tempfile
from werkzeug.utils import secure_filename
from docx import Document
import re
import json
from datetime import datetime
from functools import wraps
from config import Config as ConfigClass

# Initialize config
config = ConfigClass()

from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import RGBColor, Inches, Pt
from typing import Optional, List, Dict, Any, Tuple, Union

# Add datetime filter
def datetimeformat(value, format='%Y-%m-%d %H:%M'):
    if value is None:
        return ""
    if isinstance(value, str):
        value = datetime.fromisoformat(value)
    return value.strftime(format)

app = Flask(__name__)

# Initialize config
config = ConfigClass()
app.secret_key = 'your-secret-key-here'  # Change this to a secure secret key in production
app.jinja_env.filters['datetimeformat'] = datetimeformat

# Configure upload settings
UPLOAD_FOLDER = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'uploads')
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max file size

# Ensure upload folder exists
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

# Allowed file extensions
ALLOWED_EXTENSIONS = {'docx'}

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

class DocumentChecker:
    def __init__(self, filepath):
        # Basic file and document attributes
        self.filepath = filepath
        self.doc = Document(filepath)
        
        # Tracking and issue reporting
        self.issues = []
        self.line_issues = []
        self.current_paragraph = None
        self.current_page = 1  # Track current page number
        self.line_page_mapping = {}  # Map line numbers to page numbers
        
        # Document structure tracking
        self.current_section = None
        self.in_abstract = False
        self.in_references = False
        self.after_abstract = False
        
        # Statistics and counters
        self.total_lines = 0
        self.lines_with_issues = 0
        self.sections_checked = 0
        self.pages_skipped = 0
        
        # Headings and content tracking
        self.headings = []
        self.subheadings = []
        self.images_found = 0
        
        print("Initializing DocumentChecker...")  # Debug
        
        # Initialize rules with default values
        self.rules = {
            'font_name': 'Times New Roman',
            'heading1_size': 16,
            'heading2_size': 14,
            'normal_text_size': 12,
            'line_spacing': 1.50,
            'margin_min': 1.75,  # in inches
            'required_header': 'Project Title',
            'required_footer_left': 'Dept. name',
            'required_line_spacing': 1.50,
            'required_margin': 1.75,  # in inches
            'lines_per_page': 50,  # Average lines per page (will be adjusted)
            'pages': [],  # Track page breaks and their line numbers
            'skip_pages': 14,  # Default value
            'start_checking_from': 'abstract'  # Default value
        }
        print(f"Initial rules: {self.rules}")  # Debug
        
        # Update rules from config with default values
        try:
            print("\n=== Loading Configuration ===")
            # Use the global config instance
            global config
            if config is None:
                config = get_config()
            app_config = config
            
            # Debug: Print all attributes of the config
            print("Available config attributes:", dir(app_config))
            
            # Get config values with safe defaults
            try:
                skip_pages = app_config.skip_pages
                print(f"Found skip_pages in config: {skip_pages}")
            except AttributeError:
                skip_pages = 14
                print("Using default skip_pages: 14")
                
            try:
                start_checking = app_config.start_checking_from
                print(f"Found start_checking_from in config: {start_checking}")
            except AttributeError:
                start_checking = 'abstract'
                print("Using default start_checking_from: 'abstract'")
            
            # Ensure start_checking is a string and lowercase it
            if not isinstance(start_checking, str):
                start_checking = str(start_checking)
            start_checking = start_checking.lower()
            
            # Update rules
            self.rules.update({
                'skip_pages': skip_pages,
                'start_checking_from': start_checking
            })
            
            print(f"Final rules: skip_pages={self.rules['skip_pages']}, start_checking_from={self.rules['start_checking_from']}")
            print("=== End of Configuration Loading ===\n")
            
        except Exception as e:
            print(f"\n!!! Error loading config: {e}")
            import traceback
            traceback.print_exc()
            print("!!! Using default configuration")
            self.rules.update({
                'skip_pages': 14,
                'start_checking_from': 'abstract'
            })
            
        self.current_section = 0
        self.line_number = 0  # Track current line number
        self.current_page = 1  # Track current page number
        self.line_page_mapping = {}  # Map line numbers to page numbers
        
        # Initialize page skipping and abstract tracking
        print("\nInitializing page skipping and abstract tracking...")  # Debug
        print(f"Rules before initialization: {self.rules}")  # Debug
        
        try:
            self.skip_page_count = self.rules['skip_pages']
            print(f"Set skip_page_count to: {self.skip_page_count}")  # Debug
        except KeyError:
            self.skip_page_count = 14
            print("Using default skip_page_count: 14")  # Debug
            
        self.pages_skipped = 0
        
        try:
            start_checking = self.rules['start_checking_from']
            self.skip_until_abstract = (start_checking == 'abstract')
            print(f"Set skip_until_abstract to: {self.skip_until_abstract}")  # Debug
        except KeyError:
            self.skip_until_abstract = True
            print("Using default skip_until_abstract: True")  # Debug
    
    def check_font(self, run, line_number=None):
        """Check if font is Times New Roman"""
        if run.font and run.font.name and run.font.name != self.rules['font_name']:
            issue = f"Font should be {self.rules['font_name']}, found '{run.font.name}'"
            if line_number is not None:
                return [f"Line {line_number}: {issue}"]
            return [issue]
        return []
    
    def check_font_size(self, para, run, line_number=None):
        """Check if font size matches the style"""
        if not run.font or not run.font.size or not run.text.strip():
            return []
            
        size_pt = run.font.size.pt
        expected_size = self.rules['normal_text_size']  # Default to normal size
        
        # Determine expected size based on style
        style_name = para.style.name.lower()
        if 'heading' in style_name:
            if '1' in style_name:
                expected_size = self.rules['heading1_size']
            elif '2' in style_name:
                expected_size = self.rules['heading2_size']
        
        if abs(size_pt - expected_size) > 0.1:  # Allow for small rounding differences
            issue = f"Font size should be {expected_size}pt for '{style_name}', found {size_pt}pt"
            if line_number is not None:
                return [f"Line {line_number}: {issue}"]
            return [issue]
        return []
    
    def check_alignment(self, para, line_number=None):
        """Check if paragraph alignment is correct"""
        # Skip if no alignment set (default is left) or empty paragraph
        if not hasattr(para, 'alignment') or not para.alignment or not para.text.strip():
            return []
            
        # Get the text for context
        text = para.text.lower().strip()
        
        # Default expected alignment is JUSTIFY for normal text
        expected_alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        # Check for exceptions (headings, titles, etc.)
        if any(keyword in text for keyword in ['title', 'chapter', 'abstract', 'acknowledgment', 'appendix', 'reference']):
            expected_alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        if para.alignment != expected_alignment:
            alignment_names = {
                WD_ALIGN_PARAGRAPH.LEFT: 'left',
                WD_ALIGN_PARAGRAPH.CENTER: 'center',
                WD_ALIGN_PARAGRAPH.RIGHT: 'right',
                WD_ALIGN_PARAGRAPH.JUSTIFY: 'justify'
            }
            issue = f"Alignment should be {alignment_names.get(expected_alignment, 'justify')}, found {alignment_names.get(para.alignment, 'unknown')}"
            if line_number is not None:
                return [f"Line {line_number}: {issue}"]
            return [issue]
        return []
    
    def check_margins(self):
        try:
            sections = self.doc.sections
            for section in sections:
                try:
                    left_margin = float(str(section.left_margin).replace('Inches', '').strip())
                    right_margin = float(str(section.right_margin).replace('Inches', '').strip())
                    if left_margin < self.rules['margin_min'] or right_margin < self.rules['margin_min']:
                        self.issues.append(f"Margins should be at least {self.rules['margin_min']} inches on both sides")
                        break
                except (ValueError, AttributeError) as e:
                    self.issues.append("Could not verify margin sizes. Please check margins manually.")
                    print(f"Margin check error: {e}")
                    break
        except Exception as e:
            print(f"Error in margin checking: {e}")
            self.issues.append("Error checking document margins.")
    
    def add_issue(self, message, line_num=None, line_text=''):
        """Add an issue with line number and context"""
        if line_num is not None:
            # Convert line_num to int if it's a string that can be converted to float
            try:
                if isinstance(line_num, str):
                    line_num = int(round(float(line_num)))
                elif isinstance(line_num, float):
                    line_num = int(round(line_num))
                elif not isinstance(line_num, int):
                    # If it's not a string, float, or int, use 0 as fallback
                    line_num = 0
            except (ValueError, TypeError):
                line_num = 0
                
            self.line_issues.append({
                'line': line_num,
                'text': str(line_text)[:100] + ('...' if len(str(line_text)) > 100 else ''),
                'issue': str(message)
            })
        self.issues.append(str(message))
    
    def update_section_tracking(self, text):
        """Update section tracking based on text content"""
        text_lower = text.lower()
        
        # Check for abstract section
        if 'abstract' in text_lower and not self.after_abstract:
            self.in_abstract = True
            self.current_section = 'abstract'
        # Check for references section
        elif any(keyword in text_lower for keyword in ['references', 'bibliography']) and not self.in_references:
            self.in_references = True
            self.after_abstract = True
            self.current_section = 'references'
        # If we were in abstract and now we're not, mark after_abstract as True
        elif self.in_abstract and not self.after_abstract and text.strip():
            self.in_abstract = False
            self.after_abstract = True
            self.current_section = 'main_content'
            
    def check_image_alignment(self, para, line_num):
        """Check if images are center aligned"""
        issues = []
        
        # Check if paragraph contains an image
        for run in para.runs:
            if run._element.xpath('.//a:blip'):  # Check for images in the run
                # Check paragraph alignment
                if para.paragraph_format.alignment != WD_ALIGN_PARAGRAPH.CENTER:
                    issues.append({
                        'message': "Center align images",
                        'type': 'formatting',
                        'severity': 'high',
                        'line': line_num
                    })
                # Check if image is too large (wider than 6 inches)
                try:
                    for drawing in run._element.xpath('.//wp:inline'):
                        extents = drawing.xpath('.//a:ext', namespaces=drawing.nsmap)
                        if extents:
                            width = int(extents[0].get('cx', 0)) / 914400  # Convert EMUs to inches
                            if width > 6:  # If image is wider than 6 inches
                                issues.append({
                                    'message': f"Image is too wide ({width:.1f} inches). Resize to be 6 inches or less.",
                                    'type': 'formatting',
                                    'severity': 'medium',
                                    'line': line_num
                                })
                except Exception as e:
                    print(f"Error checking image dimensions: {e}")
                break  # Only need to check once per paragraph
                
        return issues
        
    def __init__(self, filepath):
        self.filepath = filepath
        self.doc = Document(filepath)
        self.issues = []
        self.line_issues = []
        self.current_section = None
        self.in_abstract = False
        self.in_references = False
        self.after_abstract = False
        self.total_lines = 0
        self.lines_with_issues = 0
        self.sections_checked = 0
        self.current_paragraph = None
        self.rules = {
            'font_name': 'Times New Roman',
            'heading1_size': 16,
            'heading2_size': 14,
            'normal_text_size': 12,
            'line_spacing': 1.50,
            'margin_min': 1.75,  # in inches
            'required_header': 'Project Title',
            'required_footer_left': 'Dept. name',
            'required_line_spacing': 1.50,
            'required_margin': 1.75,  # in inches
            'lines_per_page': 50,  # Average lines per page (will be adjusted)
            'pages': []  # Track page breaks and their line numbers
        }
        
        # Define expected chapter structure based on the provided outline
        self.expected_structure = {
            "ABSTRACT": {},
            "CHAPTER 1: INTRODUCTION": {
                "1.1": "GENERAL",
                "1.2": "SCOPE OF THE PROJECT",
                "1.3": "OBJECTIVE",
                "1.4": "EXISTING SYSTEM",
                "1.4.1": "EXISTING SYSTEM DISADVANTAGES",
                "1.4.2": "LITERATURE SURVEY",
                "1.5": "PROPOSED SYSTEM",
                "1.5.1": "PROPOSED SYSTEM ADVANTAGE"
            },
            "CHAPTER 2: PROJECT DESCRIPTION": {
                "2.1": "GENERAL",
                "2.2": "METHODOLOGIES",
                "2.2.1": "MODULES",
                "2.2.2": "MODULES DIAGRAMS",
                "2.3": "UI DESIGN",
                "2.3.1": "USER INTERFACE DESIGN",
                "2.3.2": "USERS",
                "2.4": "GIVEN INPUT EXPECTED OUTPUT",
                "2.5": "TECHNIQUE OR ALGORITHM USED",
                "2.5.1": "PROPOSED ALGORITHM"
            },
            "CHAPTER 3: REQUIREMENTS ENGINEERING": {
                "3.1": "GENERAL",
                "3.2": "HARDWARE REQUIREMENTS",
                "3.3": "SOFTWARE REQUIREMENTS",
                "3.4": "FUNCTIONAL REQUIREMENTS",
                "3.5": "NON-FUNCTIONAL REQUIREMENTS",
                "3.6": "DOMAIN REQUIREMENT"
            },
            "CHAPTER 4: SYSTEM DESIGN": {
                "4.1": "GENERAL",
                "4.2": "SYSTEM ARCHITECTURE",
                "4.3": "UML",
                "4.3.1": "USE CASE DIAGRAM",
                "4.3.2": "CLASS DIAGRAM",
                "4.3.3": "OBJECT DIAGRAM",
                "4.3.4": "COMPONENT DIAGRAM",
                "4.3.5": "DEPLOYMENT DIAGRAM",
                "4.3.6": "SEQUENCE DIAGRAM",
                "4.3.7": "COLLABORATION DIAGRAM",
                "4.3.8": "STATE DIAGRAM",
                "4.3.9": "ACTIVITY DIAGRAM",
                "4.4": "DATA FLOW DIAGRAM",
                "4.5": "E-R DIAGRAM",
                "4.6": "GUI DESIGN",
                "4.6.1": "COMPONENTS OF GUI",
                "4.6.2": "FEATURES OF GUI"
            },
            "CHAPTER 5: IMPLEMENTATION": {
                "5.1": "GENERAL",
                "5.2": "IMPLEMENTATION"
            },
            "CHAPTER 6: SNAPSHOTS": {
                "6.1": "GENERAL",
                "6.2": "OUTPUT SNAPSHOTS"
            },
            "CHAPTER 7: SOFTWARE TESTING": {
                "7.1": "GENERAL",
                "7.2": "DEVELOPING METHODOLOGIES",
                "7.3": "TEST STRATEGY",
                "7.3.1": "LEVELS OF TESTING",
                "7.3.2": "TYPES OF TESTING",
                "7.3.3": "TEST CASE TYPE – GUI",
                "7.3.4": "TEST DESIGN TECHNIQUES",
                "7.3.5": "TEST ENVIRONMENT",
                "7.4": "ACCEPTANCE CRITERIA",
                "7.4.1": "ACCEPTANCE TESTING",
                "7.5": "BUILD THE TEST PLAN"
            },
            "CHAPTER 8: CONCLUSION AND REFERENCES": {
                "8.1": "CONCLUSION",
                "8.2": "FUTURE ENHANCEMENT",
                "8.3": "REFERENCES"
            }
        }
        
        # Track found sections for validation
        self.found_sections = {chapter: {} for chapter in self.expected_structure}
        self.current_chapter = None
        self.missing_sections = []
        self.extra_sections = []
        
    def is_chapter_heading(self, text):
        """Check if the text is a chapter heading"""
        # Remove any extra whitespace and normalize case for comparison
        normalized_text = ' '.join(text.strip().upper().split())
        return normalized_text in self.expected_structure
    
    def is_section_heading(self, text):
        """Check if the text is a section or subsection heading"""
        if not self.current_chapter:
            return False
            
        # Check if text matches the pattern "X.Y[.Z] TITLE"
        import re
        match = re.match(r'^(\d+(?:\.\d+)*)\s+(.+)$', text.strip())
        if not match:
            return False
            
        section_num, section_text = match.groups()
        # Clean up section text: remove any trailing numbers and extra spaces
        cleaned_section_text = re.sub(r'\s*\d+\s*$', '', section_text.strip()).upper()
        cleaned_section_text = ' '.join(cleaned_section_text.split())  # Normalize spaces
        
        # Check if this section exists in the current chapter
        if section_num in self.expected_structure.get(self.current_chapter, {}):
            expected_text = self.expected_structure[self.current_chapter][section_num]
            # Compare with cleaned section text
            if cleaned_section_text == expected_text:
                # Mark this section as found
                self.found_sections[self.current_chapter][section_num] = True
                # If original text was different, add a warning
                if section_text.upper() != expected_text:
                    self.issues.append(f"Warning: Section {section_num} has extra characters. Expected: '{expected_text}', Found: '{section_text}'")
                return True
            else:
                # Found section number but text doesn't match
                self.issues.append(f"Section {section_num} has incorrect title. Expected: '{expected_text}', Found: '{section_text}'")
                return True
                
        # If we get here, it's either an extra section or a section in the wrong chapter
        for chapter, sections in self.expected_structure.items():
            if section_num in sections and cleaned_section_text == sections[section_num]:
                self.issues.append(f"Section {section_num} '{section_text}' appears to be in the wrong chapter. Expected in: {chapter}")
                return True
                
        # If we get here, it's an extra section not in our expected structure
        self.extra_sections.append(f"{section_num} {section_text}")
        return True
    
    def validate_structure(self):
        """Validate the document structure against expected headings"""
        # Check for missing chapters
        for chapter in self.expected_structure:
            if chapter not in self.found_sections or not self.found_sections[chapter]:
                self.missing_sections.append(f"Missing chapter: {chapter}")
            
            # Check for missing sections in each chapter
            if chapter in self.found_sections:
                for section_num, section_text in self.expected_structure[chapter].items():
                    if section_num not in self.found_sections[chapter]:
                        self.missing_sections.append(f"Missing section: {chapter} -> {section_num} {section_text}")
        
        # Add missing sections to issues
        if self.missing_sections:
            self.issues.append("Document structure issues found:")
            self.issues.extend(self.missing_sections)
            
        # Add extra sections to issues
        if self.extra_sections:
            self.issues.append("\nUnexpected sections found in document:")
            self.issues.extend(self.extra_sections)
    
    def check_document_structure(self):
        """Check the document structure against the expected format"""
        # Reset tracking variables
        self.issues = []
        self.found_sections = {chapter: {} for chapter in self.expected_structure}
        self.current_chapter = None
        self.missing_sections = []
        self.extra_sections = []
        
        # First pass: identify chapters and sections
        for para in self.doc.paragraphs:
            line_text = para.text.strip()
            if not line_text:
                continue
                
            # Check for chapter headings
            if self.is_chapter_heading(line_text):
                self.current_chapter = ' '.join(line_text.upper().split())
                continue
                
            # Check for section headings
            if self.current_chapter:
                self.is_section_heading(line_text)
        
        # Validate the structure
        self.validate_structure()
        
        return {
            'missing_sections': self.missing_sections,
            'extra_sections': self.extra_sections,
            'issues': self.issues
        }
        
    def check_document(self):
        """Main method to check the entire document"""
        # First check document structure
        structure_results = self.check_document_structure()
        
        # Reset tracking variables for other checks
        self.issues = []
        self.line_issues = []
        self.total_lines = 0
        self.lines_with_issues = 0
        self.sections_checked = 0
        self.current_page = 1
        self.line_page_mapping = {}
        self.current_section = None
        self.in_abstract = False
        self.in_references = False
        self.after_abstract = False
        self.pages_skipped = 0
        self.headings = []
        self.subheadings = []
        self.images_found = 0
        
        # Initialize rules if not already set
        if not hasattr(self, 'rules') or not self.rules:
            self.rules = {}
            
        # Initialize pages list in rules
        self.rules['pages'] = []
        
        # Set default values if not present
        if 'start_checking_from' not in self.rules:
            self.rules['start_checking_from'] = 'abstract'
        if 'skip_pages' not in self.rules:
            self.rules['skip_pages'] = 14
            
        # Now it's safe to access these values
        self.skip_until_abstract = (self.rules['start_checking_from'] == 'abstract')
        self.skip_page_count = self.rules['skip_pages']
        
        print(f"Will skip first {self.skip_page_count} pages")
        
        # Check document-wide settings
        self.check_margins()
        self.check_page_numbering()
        
        # First pass: count total lines and find page breaks
        total_paragraphs = sum(1 for _ in self.doc.paragraphs if _.text.strip())
        self.rules['lines_per_page'] = max(40, min(60, total_paragraphs // 10))
        
        # Track if we're still in the first 14 pages
        in_skipped_pages = True
        
        # Second pass: check content
        for para in self.doc.paragraphs:
            self.total_lines += 1
            self.current_paragraph = para
            line_text = para.text.strip()
            
            # Update page tracking first - this updates self.current_page
            page_break_found = self.update_page_break(para)
            
            # Skip all content in first 14 pages
            if in_skipped_pages:
                if page_break_found:
                    print(f"Skipping page {self.current_page}")
                    if self.current_page >= self.skip_page_count:
                        in_skipped_pages = False
                        print(f"Reached page {self.current_page}, starting checks...")
                continue
                
            # From here on, we're past the first 14 pages
            
            # Skip empty paragraphs unless they contain page breaks
            if not line_text and not page_break_found:
                continue
                
            # Check for chapter headings
            if self.is_chapter_heading(line_text):
                self.current_chapter = ' '.join(line_text.upper().split())
                continue
                
            # Check for section headings
            if self.current_chapter and self.is_section_heading(line_text):
                continue
                
            # Update section tracking
            prev_section = self.current_section
            self.update_section_tracking(line_text)
            
            # Handle abstract checking if needed
            if self.skip_until_abstract and self.in_abstract:
                self.skip_until_abstract = False
                self.after_abstract = True
                
            if prev_section != self.current_section:
                self.sections_checked += 1
            
            # Skip empty paragraphs unless they contain page breaks
            if not line_text and not page_break_found:
                continue
                
            # Initialize issues for this line
            line_has_issues = False
            line_issues = []
            
            # Check paragraph-level formatting
            alignment_issues = self.check_alignment(para, self.total_lines)
            if alignment_issues:
                line_issues.extend(alignment_issues)
                line_has_issues = True
                
            # Check for images and their alignment
            image_issues = self.check_image_alignment(para, self.total_lines)
            if image_issues:
                line_issues.extend(image_issues)
                line_has_issues = True
                self.images_found += 1
            
            # Check runs within the paragraph
            for run in para.runs:
                font_issues = self.check_font(run, self.total_lines)
                if font_issues:
                    line_issues.extend(font_issues)
                    line_has_issues = True
                
                size_issues = self.check_font_size(para, run, self.total_lines)
                if size_issues:
                    line_issues.extend(size_issues)
                    line_has_issues = True
                
                color_issues = self.check_text_color(run, self.total_lines)
                if color_issues:
                    line_issues.extend(color_issues)
                    line_has_issues = True
            
            # Add line to issues if it has any problems
            if line_has_issues or page_break_found:
                self.lines_with_issues += 1
                self.line_issues.append({
                    'line_number': self.total_lines,
                    'page_number': self.current_page,
                    'text': line_text[:200] + ('...' if len(line_text) > 200 else ''),
                    'issues': line_issues,
                    'is_page_break': page_break_found
                })
                if line_issues:  # Only add to main issues if there are actual issues
                    self.issues.extend(line_issues)
        
        # Check headers and footers
        header_footer_issues = self.check_headers_footers()
        if header_footer_issues:
            self.issues.extend(header_footer_issues)
        
        # Check page numbering sequence
        page_num_issues = self.check_page_number_sequence()
        if page_num_issues:
            self.issues.extend(page_num_issues)
        
        # Calculate total issues
        total_issues = len(self.issues) + len(self.line_issues)
        
        # Create summary
        summary = {
            'total_issues': total_issues,
            'lines_checked': self.total_lines,
            'lines_with_issues': self.lines_with_issues,
            'sections_checked': self.sections_checked,
            'compliance_score': max(0, 100 - (total_issues * 2)),  # Simple scoring (0-100)
            'heading_count': len(self.headings),
            'subheading_count': len(self.subheadings),
            'page_count': len(self.rules.get('pages', [1])) or 1
        }
        
        # Process line issues to ensure they have all required fields
        processed_line_issues = []
        for issue in self.line_issues:
            if not isinstance(issue, dict):
                continue
                
            processed_issue = {
                'line_number': issue.get('line_number', 0),
                'text': issue.get('text', ''),
                'page_number': issue.get('page_number', 1),
                'is_page_break': issue.get('is_page_break', False),
                'issues': issue.get('issues', [])
            }
            processed_line_issues.append(processed_issue)
        
        return {
            'issues': self.issues,
            'line_issues': processed_line_issues,
            'headings': self.headings,
            'subheadings': self.subheadings,
            'summary': summary
        }
        
    def check_headers_footers(self):
        """Check headers and footers after abstract"""
        issues = []
        if not self.after_abstract:
            return issues
            
        for section in self.doc.sections:
            # Check header
            if section.header:
                header_text = ' '.join([p.text for p in section.header.paragraphs])
                if self.rules['required_header'].lower() not in header_text.lower():
                    issues.append(f"Header should contain: '{self.rules['required_header']}'")
            
            # Check footer
            if section.footer:
                footer_text = ' '.join([p.text for p in section.footer.paragraphs])
                if (self.rules['required_footer_left'].lower() not in footer_text.lower() or 
                    'page' not in footer_text.lower()):
                    issues.append("Footer should contain department name and page number")
                    
        return issues
            
    def estimate_page_number(self, line_num):
        """Estimate page number based on line number and content"""
        # Convert line_num to int to ensure it's a valid dictionary key
        line_num_int = int(round(line_num))
        
        # If we've already mapped this line to a page, return that
        if line_num_int in self.line_page_mapping:
            return self.line_page_mapping[line_num_int]
            
        # Otherwise estimate based on lines per page
        estimated_page = (line_num_int // self.rules['lines_per_page']) + 1
        self.line_page_mapping[line_num_int] = estimated_page
        return estimated_page
        
    def update_page_break(self, para):
        """Check if paragraph contains a page break and update current page"""
        # Check for explicit page breaks
        if 'w:br' in para._p.xml and 'type="page"' in para._p.xml:
            self.current_page += 1
            self.rules['pages'].append({
                'page': self.current_page,
                'line': self.total_lines,
                'content': para.text[:100] + ('...' if len(para.text) > 100 else '')
            })
            return True
        return False
    
    def check_page_numbering(self):
        """Check page numbering format (Roman before abstract, numbers after)"""
        # This is a simplified check as python-docx has limited access to page numbers
        self.issues.append("Note: Page numbers are estimated. Please verify manually - Roman numerals before abstract, numbers after")
        
    def check_page_number_sequence(self):
        """Check if page numbers follow the correct sequence"""
        # This is a simplified check as detailed page number sequence checking is complex
        # and would require more sophisticated document analysis
        return []
    

    def check_lists(self):
        """Check that lists use bullets, not Roman numerals"""
        for para in self.doc.paragraphs:
            if para.style.name.startswith('List'):
                if 'List Number' in para.style.name:
                    self.issues.append("Use bullet points for lists, not numbered lists")
                    break  # Only show this warning once
    
    def check_text_color(self, run, line_number=None):
        """Check if text color is black"""
        if run.font and run.font.color and run.font.color.rgb and run.font.color.rgb != RGBColor(0, 0, 0):
            issue = "Text color should be black"
            if line_number is not None:
                return [f"Line {line_number}: {issue}"]
            return [issue]
        return []

# Admin authentication decorator
def admin_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if 'admin_logged_in' not in session:
            return redirect(url_for('admin_login'))
        return f(*args, **kwargs)
    return decorated_function

@app.route('/')
def index():
    return render_template('index.html')

# Admin routes
@app.route('/admin', methods=['GET'])
@admin_required
def admin_dashboard():
    return render_template('admin_dashboard.html', config=config)

@app.route('/admin/login', methods=['GET', 'POST'])
def admin_login():
    if request.method == 'POST':
        username = request.form.get('username')
        password = request.form.get('password')
        
        if username == config.admin_username and password == config.admin_password:
            session['admin_logged_in'] = True
            return redirect(url_for('admin_dashboard'))
        else:
            flash('Invalid username or password')
    
    return render_template('admin_login.html')

@app.route('/admin/logout')
def admin_logout():
    session.pop('admin_logged_in', None)
    return redirect(url_for('index'))

@app.route('/admin/update-settings', methods=['POST'])
@admin_required
def update_settings():
    # Update configuration
    config.skip_pages = int(request.form.get('skip_pages', 14))
    config.start_checking_from = request.form.get('start_checking_from', 'abstract')
    config.required_font = request.form.get('required_font', 'Times New Roman')
    config.required_font_size = float(request.form.get('required_font_size', 12))
    config.required_line_spacing = float(request.form.get('required_line_spacing', 1.50))
    
    # Save to config file
    config.save()
    
    flash('Settings updated successfully!')
    return redirect(url_for('admin_dashboard'))

@app.route('/admin/export-checks')
@admin_required
def export_checks():
    """Export document checks as CSV"""
    import csv
    from io import StringIO
    from flask import Response
    
    # Create CSV in memory
    si = StringIO()
    cw = csv.writer(si)
    
    # Write header
    cw.writerow(['Timestamp', 'Filename', 'Issues Found', 'IP Address', 'Page Count', 'Line Count', 'Sections Checked'])
    
    # Write data
    for check in config.document_checks:
        cw.writerow([
            check['timestamp'],
            check['filename'],
            check['issues_found'],
            check['user_ip'],
            check['metadata'].get('page_count', ''),
            check['metadata'].get('line_count', ''),
            check['metadata'].get('sections_checked', '')
        ])
    
    # Create response
    output = si.getvalue()
    return Response(
        output,
        mimetype="text/csv",
        headers={"Content-disposition": "attachment; filename=document_checks.csv"}
    )

@app.route('/admin/delete-check', methods=['POST'])
@admin_required
def delete_check():
    """Delete a specific document check"""
    try:
        data = request.get_json()
        index = data.get('index')
        
        if index is None or not isinstance(index, int) or index < 0 or index >= len(config.document_checks):
            return jsonify({'success': False, 'error': 'Invalid check index'}), 400
        
        # Remove the check at the specified index
        config.document_checks.pop(index)
        config.save()
        
        return jsonify({'success': True})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/admin/clear-all-checks', methods=['POST'])
@admin_required
def clear_all_checks():
    """Clear all document checks"""
    try:
        config.document_checks = []
        config.save()
        return jsonify({'success': True})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/upload', methods=['POST'])
def upload_file():
    # Check if the post request has the file part
    if 'file' not in request.files:
        return jsonify({'error': 'No file part in the request'}), 400
    
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'No selected file'}), 400
    
    if not (file and allowed_file(file.filename)):
        return jsonify({'error': 'Invalid file type. Please upload a .docx file'}), 400
    
    try:
        # Save the file to a temporary location
        temp_dir = tempfile.mkdtemp()
        temp_file_path = os.path.join(temp_dir, secure_filename(file.filename))
        file.save(temp_file_path)
        
        try:
            # Process the document
            checker = DocumentChecker(temp_file_path)
            issues = checker.check_document()
            
            # Prepare results
            result = {
                'filename': secure_filename(file.filename),
                'timestamp': datetime.now().isoformat(),
                'issues': issues,
                'summary': {
                    'total_issues': len(issues),
                    'lines_checked': len(checker.line_issues),
                    'lines_with_issues': sum(1 for _, data in (checker.line_issues or []) if data and 'issues' in data and data['issues']),
                    'sections_checked': len(set(getattr(checker, 'sections_checked', []))),
                    'heading_count': len(getattr(checker, 'headings', [])),
                    'subheading_count': len(getattr(checker, 'subheadings', []))
                },
                'line_issues': [(num, data) for num, data in (checker.line_issues or []) if data and 'issues' in data and data['issues']],
                'headings': getattr(checker, 'headings', []),
                'subheadings': getattr(checker, 'subheadings', [])
            }
            
            # Store minimal data in session
            session['last_result'] = {
                'filename': result['filename'],
                'timestamp': result['timestamp'],
                'summary': result['summary']
            }
            
            # Return only the result data, not the file
            return jsonify({
                'success': True,
                'result': result
            })
            
        except Exception as e:
            return jsonify({'error': f'Error processing file: {str(e)}'}), 500
            
        finally:
            # Clean up the temporary file and directory
            try:
                if os.path.exists(temp_file_path):
                    os.unlink(temp_file_path)
                if os.path.exists(temp_dir):
                    os.rmdir(temp_dir)
            except Exception as e:
                print(f"Error removing temporary files: {e}")
                
    except Exception as e:
        # Clean up in case of error
        if 'temp_dir' in locals() and os.path.exists(temp_dir):
            try:
                if 'temp_file_path' in locals() and os.path.exists(temp_file_path):
                    os.unlink(temp_file_path)
                os.rmdir(temp_dir)
            except Exception as cleanup_error:
                print(f"Error during cleanup: {cleanup_error}")
        return jsonify({'error': f'Error handling file upload: {str(e)}'}), 500

# Vercel handler
app = app

def vercel_handler(event, context):
    with app.app_context():
        response = app.full_dispatch_request()
        return {
            'statusCode': response.status_code,
            'headers': dict(response.headers),
            'body': response.get_data(as_text=True)
        }

# For local development
if __name__ == '__main__':
    app.run(debug=True)
